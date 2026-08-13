# -*- coding: utf-8 -*-
"""把 Office 文档转成纯文本，供 git diff 做内容级比对（textconv 驱动）。

装法（项目里做一次）：
    mkdir -p .gittools && cp office2txt.py .gittools/
    git config diff.office.textconv "python .gittools/office2txt.py"
    printf '*.docx binary diff=office\\n*.xlsx binary diff=office\\n' > .gitattributes

之后 git diff / git log -p 就能显示逐段文字、逐单元格的差异，而不是
"Bin 35533 -> 35565 bytes"。

单独使用：python office2txt.py <文件>
"""
import os
import sys


def docx_text(path):
    from docx import Document
    doc = Document(path)
    out = []
    # 刻意不输出段落序号：插入/删除一段会让其后所有序号平移，
    # 导致 diff 满屏"改动"，淹没真正的实质修改。
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                out.append("[表] " + " | ".join(cells))
    # 脚注：正文之外的独立部件，很容易被悄悄改掉而无人察觉
    try:
        from lxml import etree
        W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for part in doc.part.package.parts:
            if str(part.partname) != "/word/footnotes.xml":
                continue
            root = etree.fromstring(part._blob)
            for fn in root.findall(W + "footnote"):
                if fn.get(W + "type"):      # separator 之类，跳过
                    continue
                t = "".join(fn.itertext()).strip()
                if t:
                    out.append("[脚注] " + t)   # 同样不输出 id，Office 保存时会重排
    except Exception as e:
        out.append("[脚注解析失败] %s" % e)
    return "\n".join(out)


def xlsx_text(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=False)   # 保留公式，能看出公式被改成常量
    out = []
    for ws in wb.worksheets:
        out.append("### 工作表: %s" % ws.title)
        for row in ws.iter_rows():
            vals = ["%s=%s" % (c.coordinate, str(c.value).replace("\n", " "))
                    for c in row if c.value is not None]
            if vals:
                out.append("  " + " | ".join(vals))
    return "\n".join(out)


def main():
    path = sys.argv[1]
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            print(docx_text(path))
        elif ext == ".xlsx":
            print(xlsx_text(path))
        else:
            # .doc/.xls/.ppt 是旧式二进制，无法直接解析；输出指纹，
            # 至少能判断"变没变"，避免 diff 报错
            import hashlib
            print("[旧式二进制 %s] md5=%s size=%d"
                  % (ext, hashlib.md5(open(path, "rb").read()).hexdigest(),
                     os.path.getsize(path)))
    except Exception as e:
        print("[解析失败 %s] %s" % (path, e))


if __name__ == "__main__":
    main()
